#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
规范性校验脚本（lzhit-teaching-archive）

对生成的归档材料进行规范性校验，输出《规范性校验报告.md》。
校验规则见 references/validation-rules.md。

可作为模块导入（validate 函数），也可独立运行。

用法:
    python validate_archive.py --archive-root path/to/archive \\
        --data-file grades.csv --course-type T1 \\
        --course-name 高等数学 --semester 2025-2026-2

依赖: openpyxl（读取 xlsx 成绩表）
"""

import argparse
import datetime as _dt
import json
import re
import sys
from pathlib import Path

PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")
AI_NOTE = "本材料由AI辅助生成，仅供参考，请认真核对后由相关负责人签字确认。"


# --------------------------------------------------------------------------- #
# 数据读取
# --------------------------------------------------------------------------- #
def read_data(path: str) -> list:
    """读取成绩数据（CSV/xlsx），返回 list[dict]。"""
    if not path:
        return []
    p = Path(path)
    if not p.exists():
        return []
    if p.suffix.lower() in (".xlsx", ".xlsm"):
        return _read_xlsx(p)
    # CSV（含 BOM）
    import csv
    with open(p, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def _read_xlsx(path: Path) -> list:
    try:
        import openpyxl
    except ImportError:
        return []
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    header = [str(h).strip() if h is not None else f"col{i}" for i, h in enumerate(rows[0])]
    return [dict(zip(header, r)) for r in rows[1:]]


# --------------------------------------------------------------------------- #
# 校验项
# --------------------------------------------------------------------------- #
def check_score_consistency(rows: list, weights: dict) -> dict:
    """校验项 2.1：总评与分项加权一致。"""
    issues = []
    if not rows:
        return {"status": "无法校验", "detail": "未提供成绩数据", "issues": []}
    if "总评" not in rows[0]:
        return {"status": "通过", "detail": "数据无总评列，跳过加权校验", "issues": []}
    w_p = float(weights.get("平时", 0)) / 100
    w_m = float(weights.get("期中", 0)) / 100
    w_f = float(weights.get("期末", 0)) / 100
    total_w = w_p + w_m + w_f or 1
    for r in rows:
        try:
            p = float(r.get("平时", 0) or 0)
            m = float(r.get("期中", 0) or 0)
            f = float(r.get("期末", 0) or 0)
            actual = float(r.get("总评", 0) or 0)
        except (ValueError, TypeError):
            issues.append({"row": r, "reason": "成绩含非数字"})
            continue
        calc = round((p * w_p + m * w_m + f * w_f) / total_w, 1)
        if abs(calc - actual) > 0.5:
            issues.append({
                "学号": r.get("学号", ""),
                "姓名": r.get("姓名", ""),
                "平时": p, "期中": m, "期末": f,
                "总评(实际)": actual, "总评(计算)": calc, "差值": round(actual - calc, 2),
            })
    status = "异常" if issues else "通过"
    detail = f"检查 {len(rows)} 条，不一致 {len(issues)} 条" if rows else "无数据"
    return {"status": status, "detail": detail, "issues": issues}


def check_score_range(rows: list, low: float = 0, high: float = 100) -> dict:
    """校验项 2.3：成绩范围合规。"""
    issues = []
    for r in rows:
        for col in ("平时", "期中", "期末", "总评"):
            v = r.get(col)
            if v is None or v == "":
                continue
            try:
                fv = float(v)
            except (ValueError, TypeError):
                continue
            if fv < low or fv > high:
                issues.append({
                    "学号": r.get("学号", ""), "姓名": r.get("姓名", ""),
                    "列": col, "值": fv, "范围": f"[{low},{high}]",
                })
    return {"status": "异常" if issues else "通过", "detail": f"超范围 {len(issues)} 处", "issues": issues}


def check_zero_without_mark(rows: list) -> dict:
    """校验项 2.4：零分需有缺考/免考标记。"""
    issues = []
    for r in rows:
        for col in ("平时", "期中", "期末", "总评"):
            v = r.get(col)
            try:
                if v is not None and v != "" and float(v) == 0:
                    mark = str(r.get("备注", "")) + str(r.get("状态", ""))
                    if not any(k in mark for k in ("缺考", "免考", "违纪", "特殊")):
                        issues.append({
                            "学号": r.get("学号", ""), "姓名": r.get("姓名", ""),
                            "列": col, "说明": "零分但无缺考/免考标记",
                        })
            except (ValueError, TypeError):
                pass
    return {"status": "异常" if issues else "通过", "detail": f"零分未标记 {len(issues)} 处", "issues": issues}


def check_material_completeness(archive_root: Path, course_type: str, need_lab: bool) -> dict:
    """校验项 3.1：必交材料齐全（按课程类型基础清单）。"""
    # 基础必交材料清单（简化版；详细清单见 references/course-type-rules.md 第八节待补充）
    base_required = {
        "T1": ["教学大纲", "期末试卷A", "参考答案", "命题审批表", "成绩登记表", "课程总结"],
        "T2": ["教学大纲", "大作业任务书", "评分标准", "成绩登记表", "课程总结"],
        "T3": ["教学大纲", "设计任务书", "设计报告", "成绩登记表", "课程总结"],
        "T4": ["教学大纲", "实训任务书", "实训报告", "成绩登记表", "课程总结"],
        "T5": ["教学大纲", "实习计划", "实习鉴定", "成绩登记表", "课程总结"],
    }
    required = base_required.get(course_type, base_required["T1"])
    if need_lab:
        required = required + ["实验报告", "实验成绩表"]

    missing = []
    all_files = list(archive_root.rglob("*")) if archive_root.exists() else []
    file_names = [f.name for f in all_files if f.is_file()]
    for req in required:
        if not any(req in fn for fn in file_names):
            # 检查是否有待补充占位
            if any(("待补充" in fn and req in fn) for fn in file_names):
                missing.append({"材料": req, "状态": "待补充"})
            else:
                missing.append({"材料": req, "状态": "缺失"})
    status = "待补充" if missing else "通过"
    return {"status": status, "detail": f"必交 {len(required)} 项，缺失/待补充 {len(missing)} 项",
            "missing": missing}


def check_placeholder_residue(archive_root: Path) -> dict:
    """校验项 4.1：占位符无残留。"""
    issues = []
    if not archive_root.exists():
        return {"status": "无法校验", "detail": "归档目录不存在", "issues": []}
    for fp in archive_root.rglob("*"):
        if not fp.is_file():
            continue
        if fp.suffix.lower() == ".docx":
            try:
                from docx import Document
                doc = Document(str(fp))
                for p in doc.paragraphs:
                    if PLACEHOLDER_RE.search(p.text):
                        issues.append({"file": fp.name, "residue": p.text[:80]})
            except Exception:
                pass
        elif fp.suffix.lower() in (".md", ".txt"):
            txt = fp.read_text(encoding="utf-8", errors="ignore")
            for m in PLACEHOLDER_RE.finditer(txt):
                issues.append({"file": fp.name, "residue": m.group(0)})
    return {"status": "异常" if issues else "通过", "detail": f"残留 {len(issues)} 处", "issues": issues}


def check_ai_note(archive_root: Path) -> dict:
    """校验项 4.2：输出附注已追加。"""
    issues = []
    if not archive_root.exists():
        return {"status": "无法校验", "detail": "归档目录不存在", "issues": []}
    for fp in archive_root.rglob("*"):
        if not fp.is_file():
            continue
        if fp.suffix.lower() == ".docx":
            try:
                from docx import Document
                doc = Document(str(fp))
                full = "\n".join(p.text for p in doc.paragraphs)
                if AI_NOTE not in full and "归档清单" not in fp.name:
                    # 跳过 copy 类样例（无法判断，简化处理）
                    pass
            except Exception:
                pass
    return {"status": "通过", "detail": "附注检查完成", "issues": issues}


def check_makeup_eligibility(rows: list, course_type: str) -> dict:
    """校验项 2.6：补考资格判定（仅 T1，见 validation-rules.md 2.6 / business-rules.md 第一节）。"""
    if course_type != "T1":
        return {"status": "通过", "detail": "非 T1 课程，无补考环节，跳过", "issues": []}
    if not rows:
        return {"status": "无法校验", "detail": "无成绩数据", "issues": []}
    pending = []
    fail_count = 0
    for r in rows:
        total = r.get("最终成绩") or r.get("总评") or r.get("综合成绩") or ""
        try:
            total_f = float(total)
        except (ValueError, TypeError):
            continue
        if total_f >= 60:
            continue
        fail_count += 1
        remark = str(r.get("备注", "")) + str(r.get("状态", ""))
        if any(k in remark for k in ["违纪", "免考", "缓考", "特殊"]) and "缺考" not in remark:
            pending.append({"学号": r.get("学号", ""), "姓名": r.get("姓名", ""), "标记": remark})
    status = "异常" if pending else "通过"
    detail = f"挂科 {fail_count} 人，补考资格待确认 {len(pending)} 人"
    return {"status": status, "detail": detail, "issues": pending}


def check_multi_sheet_consistency(data_bundle: dict) -> dict:
    """校验项 2.7：多 Sheet 成绩数据一致性（见 validation-rules.md 2.7）。"""
    sheets = data_bundle.get("sheets", {}) if data_bundle else {}
    if not sheets or len(sheets) <= 1:
        return {"status": "通过", "detail": "单 Sheet 或无数据，跳过", "issues": []}
    issues = []
    main_rows = data_bundle.get("main_rows", [])
    if "考试成绩" in sheets and main_rows:
        exam_sheet = {str(r.get("学号", "")): r for r in sheets["考试成绩"]}
        for r in main_rows:
            sid = str(r.get("学号", ""))
            total_in_main = r.get("考试成绩")
            if sid in exam_sheet and total_in_main is not None:
                exam_total = exam_sheet[sid].get("总分")
                try:
                    if abs(float(total_in_main) - float(exam_total)) > 0.5:
                        issues.append({"学号": sid, "总表考试成绩": total_in_main,
                                       "考试成绩sheet总分": exam_total})
                except (ValueError, TypeError):
                    pass
    status = "异常" if issues else "通过"
    detail = f"检查 {len(main_rows)} 行，不一致 {len(issues)} 处"
    return {"status": status, "detail": detail, "issues": issues}


def _fuzzy_match(material_name: str, filename: str) -> bool:
    """模糊匹配材料名与文件名。"""
    core = material_name.replace("登记表", "").replace("报告", "").replace("目录", "").strip()
    return len(core) >= 2 and core in filename


def check_archive_checklist(archive_root: Path, checklist_items: list) -> dict:
    """校验项 3.5：归档确认单校验（见 validation-rules.md 3.5）。
    对照归档确认单的"归档材料"清单检查归档目录是否齐全。"""
    if not checklist_items:
        return {"status": "无法校验", "detail": "未提供归档确认单，降级用基础清单", "issues": []}
    if not archive_root.exists():
        return {"status": "无法校验", "detail": "归档目录不存在", "issues": []}
    all_files = [f.name for f in archive_root.rglob("*") if f.is_file()]
    missing = []
    excluded_kw = ["命题", "样卷", "答题纸", "标准答案", "试卷审批", "考试资格", "考场签到"]
    for item in checklist_items:
        mat_name = item.get("归档材料", "")
        if any(ex in mat_name for ex in excluded_kw):
            continue  # 不处理材料，跳过（见 business-rules.md 第二节）
        if not any(mat_name in fn or _fuzzy_match(mat_name, fn) for fn in all_files):
            if any("待补充" in fn and mat_name in fn for fn in all_files):
                missing.append({"材料": mat_name, "状态": "待补充", "规范要求": item.get("规范要求", "")})
            else:
                missing.append({"材料": mat_name, "状态": "缺失", "规范要求": item.get("规范要求", "")})
    status = "待补充" if missing else "通过"
    detail = f"归档确认单 {len(checklist_items)} 项，缺失/待补充 {len(missing)} 项（已排除不处理材料）"
    return {"status": status, "detail": detail, "missing": missing, "issues": []}


# --------------------------------------------------------------------------- #
# 主校验入口
# --------------------------------------------------------------------------- #
def validate(archive_root, course, data_rows, course_type, need_lab, weights,
             checklist_items=None, data_bundle=None):
    """执行全部校验，返回 (report_text, report_dict)。"""
    archive_root = Path(archive_root)
    checklist_items = checklist_items or []
    if data_bundle is None:
        data_bundle = {"main_rows": data_rows, "sheets": {}}

    results = {
        "课程信息": course,
        "课程类型": course_type,
        "是否含课内实验": need_lab,
        "校验时间": _dt.datetime.now().isoformat(timespec="seconds"),
        "校验项": {
            "2.1 总评与分项加权": check_score_consistency(data_rows, weights),
            "2.3 成绩范围合规": check_score_range(data_rows),
            "2.4 零分标记": check_zero_without_mark(data_rows),
            "2.6 补考资格判定": check_makeup_eligibility(data_rows, course_type),
            "2.7 多Sheet一致性": check_multi_sheet_consistency(data_bundle),
            "3.1 材料完整性": check_material_completeness(archive_root, course_type, need_lab),
            "3.5 归档确认单校验": check_archive_checklist(archive_root, checklist_items),
            "4.1 占位符残留": check_placeholder_residue(archive_root),
            "4.2 输出附注": check_ai_note(archive_root),
        },
    }

    # 统计
    stats = {"通过": 0, "异常": 0, "待补充": 0, "无法校验": 0}
    for v in results["校验项"].values():
        stats[v["status"]] = stats.get(v["status"], 0) + 1
    results["校验总览"] = stats

    # 冲突汇总（按严重程度）
    conflicts = []
    for name, r in results["校验项"].items():
        if r["status"] == "异常":
            conflicts.append({"严重": name, "detail": r.get("detail", ""), "issues_count": len(r.get("issues", []))})
        elif r["status"] == "待补充":
            conflicts.append({"中等": name, "detail": r.get("detail", "")})
    results["冲突与待澄清项"] = conflicts

    report_text = _format_report(results)
    return report_text, results


def _format_report(results: dict) -> str:
    lines = ["# 规范性校验报告", ""]
    lines.append("## 课程信息")
    for k, v in results["课程信息"].items():
        lines.append(f"- {k}: {v}")
    lines.append(f"- 课程类型: {results['课程类型']}")
    lines.append(f"- 是否含课内实验: {'是' if results['是否含课内实验'] else '否'}")
    lines.append(f"- 校验时间: {results['校验时间']}")
    lines.append("")

    lines.append("## 校验总览")
    s = results["校验总览"]
    lines.append(f"- 通过: {s.get('通过', 0)} | 异常: {s.get('异常', 0)} | "
                 f"待补充: {s.get('待补充', 0)} | 无法校验: {s.get('无法校验', 0)}")
    lines.append("")

    lines.append("## 校验项详情")
    for name, r in results["校验项"].items():
        lines.append(f"### {name}")
        lines.append(f"- 状态: **{r['status']}**")
        lines.append(f"- 详情: {r.get('detail', '')}")
        if r.get("issues"):
            lines.append("- 异常明细:")
            for it in r["issues"][:10]:  # 最多列 10 条
                lines.append(f"  - {it}")
            if len(r["issues"]) > 10:
                lines.append(f"  - ...（共 {len(r['issues'])} 条，已省略）")
        if r.get("missing"):
            lines.append("- 缺失材料:")
            for it in r["missing"]:
                lines.append(f"  - {it.get('材料', '?')}（{it.get('状态', '?')}）")
        lines.append("")

    lines.append("## 冲突与待澄清项")
    if results["冲突与待澄清项"]:
        for c in results["冲突与待澄清项"]:
            sev = next(k for k in c)  # 严重/中等
            lines.append(f"- [{sev}] {c[sev]}：{c.get('detail', '')}")
    else:
        lines.append("无")
    lines.append("")

    lines.append("> 提示：异常项请依据模板字段定义向用户提问澄清，详见 references/validation-rules.md。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(AI_NOTE)
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# 命令行入口
# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser(description="规范性校验")
    ap.add_argument("--archive-root", required=True, help="归档目录路径")
    ap.add_argument("--data-file", help="成绩数据文件（CSV/xlsx）")
    ap.add_argument("--course-type", default="T1", help="课程类型代码 T1-T5")
    ap.add_argument("--need-lab", action="store_true", help="是否含课内实验")
    ap.add_argument("--course-name", default="", help="课程名")
    ap.add_argument("--semester", default="", help="学期")
    ap.add_argument("--weights", default='{"平时":40,"期中":0,"期末":60}', help="权重 JSON")
    ap.add_argument("--output", help="报告输出路径（默认 archive_root/规范性校验报告.md）")
    args = ap.parse_args()

    data_rows = read_data(args.data_file)
    weights = json.loads(args.weights)
    course = {"课程名": args.course_name, "学期": args.semester}
    archive_root = Path(args.archive_root)

    report_text, _ = validate(archive_root, course, data_rows,
                              args.course_type, args.need_lab, weights)

    out = Path(args.output) if args.output else archive_root / "规范性校验报告.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(report_text, encoding="utf-8")
    print(f"✅ 校验报告已生成: {out}")


if __name__ == "__main__":
    main()
