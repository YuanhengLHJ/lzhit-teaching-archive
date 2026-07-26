#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
lzhit-teaching-archive 核心引擎

完整流程：
1. 读取课程教学大纲 → 识别课程类型（T1-T5）+ 是否含课内实验（L1）
2. 在模板根目录下匹配对应课程类型的子文件夹
3. 加载该课程类型的归档规范（材料清单）
4. 填充模板占位符，生成 docx/xlsx 材料
5. 由教学过程数据生成成绩表、考勤表
6. 若含课内实验，额外生成课内实验归档材料
7. 在每份生成材料末尾追加 AI 辅助生成声明
8. 调用 validate_archive 进行规范性校验
9. 调用 expert_report 生成教学专家分析报告
10. 产出归档清单 + ZIP 压缩包

用法:
    python build_archive.py --config archive_config.json
    python build_archive.py --course-name 高等数学 --semester 2025-2026-2 \\
        --teacher 张老师 --class 计科2301 \\
        --syllabus syllabus.docx --data-file grades.csv \\
        --templates-dir ./templates --output-dir ./output

依赖: openpyxl, python-docx (pip install openpyxl python-docx)
"""

import argparse
import csv
import datetime as _dt
import json
import re
import shutil
import sys
import zipfile
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

import identify_course_type as ict       # noqa: E402
import validate_archive as va            # noqa: E402
import expert_report as er                # noqa: E402

SKILL_DIR = SCRIPT_DIR.parent
DEFAULT_TEMPLATES = SKILL_DIR / "assets" / "templates"

PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")
AI_NOTE = "本材料由AI辅助生成，仅供参考，请认真核对后由相关负责人签字确认。"

# --------------------------------------------------------------------------- #
# 默认归档规范（按课程类型）。详细清单见 references/course-type-rules.md 第八节（部分待补充）
# --------------------------------------------------------------------------- #
DEFAULT_CONVENTION = {
    "T1": {  # 理论课（考试）
        "categories": [
            {"folder": "01_教学文件", "docs": [
                {"name": "教学大纲", "file": "教学大纲_{{课程名}}_{{学期}}.docx", "template": "教学大纲.docx", "kind": "template"},
                {"name": "教学日历", "file": "教学日历_{{课程名}}_{{学期}}.docx", "template": "教学日历.docx", "kind": "template"},
            ]},
            {"folder": "02_考核材料", "docs": [
                {"name": "期末试卷A", "file": "期末试卷A_{{课程名}}_{{学期}}.docx", "template": "期末试卷A.docx", "kind": "template"},
                {"name": "期末试卷B", "file": "期末试卷B_{{课程名}}_{{学期}}.docx", "template": "期末试卷B.docx", "kind": "template"},
                {"name": "参考答案与评分标准", "file": "答案评分标准_{{课程名}}_{{学期}}.docx", "template": "参考答案与评分标准.docx", "kind": "template"},
                {"name": "命题审批表", "file": "命题审批表_{{课程名}}_{{学期}}.docx", "template": "命题审批表.docx", "kind": "template"},
                {"name": "试卷分析报告", "file": "试卷分析报告_{{课程名}}_{{学期}}.docx", "template": "试卷分析报告.docx", "kind": "template"},
            ]},
            {"folder": "03_成绩与考勤", "docs": [
                {"name": "成绩登记表", "file": "成绩登记表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
                {"name": "平时成绩表", "file": "平时成绩表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
                {"name": "考勤表", "file": "考勤表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
            ]},
            {"folder": "04_学生成果", "docs": [
                {"name": "学生作业样例", "file": "学生作业样例_{{课程名}}_{{学期}}.pdf", "template": "学生作业样例.pdf", "kind": "copy"},
            ]},
            {"folder": "06_教学总结", "docs": [
                {"name": "课程总结", "file": "课程总结_{{课程名}}_{{学期}}.docx", "template": "课程总结.docx", "kind": "template"},
                {"name": "教学反思", "file": "教学反思_{{课程名}}_{{学期}}.docx", "template": "教学反思.docx", "kind": "template"},
            ]},
        ],
    },
    "T2": {  # 理论课（考查）
        "categories": [
            {"folder": "01_教学文件", "docs": [
                {"name": "教学大纲", "file": "教学大纲_{{课程名}}_{{学期}}.docx", "template": "教学大纲.docx", "kind": "template"},
                {"name": "教学日历", "file": "教学日历_{{课程名}}_{{学期}}.docx", "template": "教学日历.docx", "kind": "template"},
            ]},
            {"folder": "02_考核材料", "docs": [
                {"name": "大作业任务书", "file": "大作业任务书_{{课程名}}_{{学期}}.docx", "template": "大作业任务书.docx", "kind": "template"},
                {"name": "大作业评分标准", "file": "大作业评分标准_{{课程名}}_{{学期}}.docx", "template": "大作业评分标准.docx", "kind": "template"},
                {"name": "考核分析表", "file": "考核分析表_{{课程名}}_{{学期}}.docx", "template": "考核分析表.docx", "kind": "template"},
            ]},
            {"folder": "03_成绩与考勤", "docs": [
                {"name": "成绩登记表", "file": "成绩登记表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
                {"name": "平时成绩表", "file": "平时成绩表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
                {"name": "考勤表", "file": "考勤表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
            ]},
            {"folder": "04_学生成果", "docs": [
                {"name": "大作业样例", "file": "大作业样例_{{课程名}}_{{学期}}.pdf", "template": "大作业样例.pdf", "kind": "copy"},
            ]},
            {"folder": "06_教学总结", "docs": [
                {"name": "课程总结", "file": "课程总结_{{课程名}}_{{学期}}.docx", "template": "课程总结.docx", "kind": "template"},
                {"name": "教学反思", "file": "教学反思_{{课程名}}_{{学期}}.docx", "template": "教学反思.docx", "kind": "template"},
            ]},
        ],
    },
    "T3": {  # 课程设计（基础框架，详细待补充）
        "categories": [
            {"folder": "01_教学文件", "docs": [
                {"name": "教学大纲", "file": "教学大纲_{{课程名}}_{{学期}}.docx", "template": "教学大纲.docx", "kind": "template"},
            ]},
            {"folder": "02_考核材料", "docs": [
                {"name": "设计任务书", "file": "设计任务书_{{课程名}}_{{学期}}.docx", "template": "设计任务书.docx", "kind": "template"},
                {"name": "设计报告模板", "file": "设计报告模板_{{课程名}}_{{学期}}.docx", "template": "设计报告模板.docx", "kind": "template"},
                {"name": "答辩记录表", "file": "答辩记录表_{{课程名}}_{{学期}}.docx", "template": "答辩记录表.docx", "kind": "template"},
            ]},
            {"folder": "03_成绩与考勤", "docs": [
                {"name": "成绩登记表", "file": "成绩登记表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
            ]},
            {"folder": "06_教学总结", "docs": [
                {"name": "课程总结", "file": "课程总结_{{课程名}}_{{学期}}.docx", "template": "课程总结.docx", "kind": "template"},
            ]},
        ],
    },
    "T4": {  # 实训（基础框架，详细待补充）
        "categories": [
            {"folder": "01_教学文件", "docs": [
                {"name": "教学大纲", "file": "教学大纲_{{课程名}}_{{学期}}.docx", "template": "教学大纲.docx", "kind": "template"},
            ]},
            {"folder": "02_考核材料", "docs": [
                {"name": "实训任务书", "file": "实训任务书_{{课程名}}_{{学期}}.docx", "template": "实训任务书.docx", "kind": "template"},
                {"name": "实训报告模板", "file": "实训报告模板_{{课程名}}_{{学期}}.docx", "template": "实训报告模板.docx", "kind": "template"},
            ]},
            {"folder": "03_成绩与考勤", "docs": [
                {"name": "成绩登记表", "file": "成绩登记表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
            ]},
            {"folder": "06_教学总结", "docs": [
                {"name": "课程总结", "file": "课程总结_{{课程名}}_{{学期}}.docx", "template": "课程总结.docx", "kind": "template"},
            ]},
        ],
    },
    "T5": {  # 实习（基础框架，详细待补充）
        "categories": [
            {"folder": "01_教学文件", "docs": [
                {"name": "教学大纲", "file": "教学大纲_{{课程名}}_{{学期}}.docx", "template": "教学大纲.docx", "kind": "template"},
            ]},
            {"folder": "02_考核材料", "docs": [
                {"name": "实习计划", "file": "实习计划_{{课程名}}_{{学期}}.docx", "template": "实习计划.docx", "kind": "template"},
                {"name": "实习鉴定表", "file": "实习鉴定表_{{课程名}}_{{学期}}.docx", "template": "实习鉴定表.docx", "kind": "template"},
                {"name": "实习日志模板", "file": "实习日志模板_{{课程名}}_{{学期}}.docx", "template": "实习日志模板.docx", "kind": "template"},
            ]},
            {"folder": "03_成绩与考勤", "docs": [
                {"name": "成绩登记表", "file": "成绩登记表_{{课程名}}_{{学期}}.xlsx", "kind": "data"},
            ]},
            {"folder": "06_教学总结", "docs": [
                {"name": "课程总结", "file": "课程总结_{{课程名}}_{{学期}}.docx", "template": "课程总结.docx", "kind": "template"},
            ]},
        ],
    },
}

# 课内实验附加归档材料（L1）
LAB_CONVENTION = {
    "categories": [
        {"folder": "05_课内实验", "docs": [
            {"name": "实验报告模板", "file": "实验报告模板_{{课程名}}_{{学期}}.docx", "template": "实验报告模板.docx", "kind": "template"},
            {"name": "实验成绩表", "file": "实验成绩表_{{课程名}}_{{学期}}.xlsx", "template": "实验成绩表.xlsx", "kind": "template"},
            {"name": "实验指导书", "file": "实验指导书_{{课程名}}_{{学期}}.docx", "template": "实验指导书.docx", "kind": "template"},
        ]},
    ]
}


# --------------------------------------------------------------------------- #
# 业务边界与动态模板发现（详见 references/business-rules.md、template-discovery.md）
# --------------------------------------------------------------------------- #
# T1 理论课（考试）不生成的材料（命题/阅卷/考务环节，由教师/教务秘书处理）
EXCLUDED_MATERIALS_T1 = {
    "命题计划表", "样卷", "答题纸", "标准答案及评分细则", "标准答案",
    "试卷审批表", "考试资格名单", "考场签到表", "学生答题纸", "答题纸封面",
}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp"}


def is_image_or_temp(path: Path) -> bool:
    """判断是否为图片示例或 Office 临时锁文件（应跳过，见 business-rules.md 4.2/4.3）。"""
    if path.name.startswith("~$"):
        return True
    return path.suffix.lower() in IMAGE_EXTS


def discover_templates(templates_root: Path, type_info: dict) -> dict:
    """动态扫描模板根目录，发现课程类型对应的子目录及角色文件。
    不写死任何模板内容，模板变化时自动适配。详见 template-discovery.md。
    返回 {matched_dir, lab_dir, makeup_dir, checklist_path, requirement_path,
          template_subdir, attachments, images}"""
    result = {k: None for k in ['matched_dir', 'lab_dir', 'makeup_dir',
                                 'checklist_path', 'requirement_path',
                                 'template_subdir']}
    result['attachments'] = []
    result['images'] = []
    if not templates_root or not templates_root.exists():
        return result

    type_code = type_info.get("course_type_code", "")
    for d in templates_root.iterdir():
        if not d.is_dir():
            continue
        name = d.name
        if type_code == "theory_exam" and "考试" in name and "理论课" in name:
            result['matched_dir'] = d
        elif type_code == "theory_assess" and "考查" in name and "理论课" in name:
            result['matched_dir'] = d
        elif type_code == "course_design" and "课程设计" in name:
            result['matched_dir'] = d
        elif type_code == "training" and "实训" in name:
            result['matched_dir'] = d
        elif type_code == "internship" and "实习" in name:
            result['matched_dir'] = d
        if any(kw in name for kw in ["课内实验", "课内上机", "课内实践"]):
            result['lab_dir'] = d
        if "补考" in name:
            result['makeup_dir'] = d

    main_dir = result['matched_dir'] or templates_root
    for fp in main_dir.iterdir():
        if not fp.is_file():
            continue
        if is_image_or_temp(fp):
            result['images'].append(fp)
            continue
        fn = fp.name
        if "归档确认单" in fn and fp.suffix.lower() == ".xlsx":
            # 优先选"正常班级"+"正考"版本
            if result['checklist_path'] is None or ("正常" in fn and "正考" in fn):
                result['checklist_path'] = fp
        elif ("要求" in fn and "提交" in fn) or fn.startswith("！！"):
            result['requirement_path'] = fp
    # 模板子目录（"模板（XXX）"）
    for d in main_dir.iterdir():
        if d.is_dir() and d.name.startswith("模板"):
            result['template_subdir'] = d
            break
    # 收集附件模板（跳过图片/临时文件）
    tsub = result['template_subdir'] or main_dir
    if tsub.exists():
        for fp in tsub.iterdir():
            if fp.is_file():
                if is_image_or_temp(fp):
                    result['images'].append(fp)
                else:
                    result['attachments'].append(fp)
    return result


def parse_archive_checklist(checklist_path: Path) -> list:
    """解析归档确认单，返回材料清单 [{序号, 归档材料, 规范要求}]。
    动态识别列位置（序号可能在第1或第2列），作为校验依据（见 validation-rules.md 3.5）。"""
    if not checklist_path or not checklist_path.exists():
        return []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(checklist_path), data_only=True)
        items = []
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            # 找表头行：同时含"序号"和"归档材料"的行
            header_idx, seq_col, mat_col, req_col = -1, -1, -1, -1
            for i, row in enumerate(rows):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if "序号" in cells and "归档材料" in cells:
                    header_idx = i
                    seq_col = cells.index("序号")
                    mat_col = cells.index("归档材料")
                    req_col = cells.index("规范要求") if "规范要求" in cells else -1
                    break
            if header_idx < 0:
                continue
            # 读取数据行：序号列为数字
            for row in rows[header_idx + 1:]:
                cells = [str(c).strip() if c is not None else "" for c in row]
                if (seq_col < len(cells) and cells[seq_col].isdigit()
                        and mat_col < len(cells) and cells[mat_col]):
                    items.append({
                        "序号": cells[seq_col],
                        "归档材料": cells[mat_col],
                        "规范要求": cells[req_col] if req_col >= 0 and req_col < len(cells) else "",
                    })
        return items
    except Exception:
        return []


def read_data_multi_sheet(data_file: str) -> dict:
    """读取多 Sheet 成绩数据。返回 {main_rows, sheets}。
    main_rows 用于补考判定/归档/分析；sheets 保留全部 sheet 供多sheet校验。
    详见 business-rules.md 第五节。"""
    if not data_file:
        return {"main_rows": [], "sheets": {}}
    p = Path(data_file)
    if not p.exists():
        return {"main_rows": [], "sheets": {}}
    if p.suffix.lower() in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(str(p), data_only=True)
        sheets = {}
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                sheets[ws.title] = []
                continue
            # 支持双行表头：第一行含 None 时合并第二行
            # 注意：合并单元格分组标题场景下，第二行子项是实际列名，非None时覆盖第一行
            header = list(rows[0])
            data_start = 1
            if len(rows) > 1 and any(h is None for h in header):
                row2 = rows[1]
                for i in range(min(len(header), len(row2))):
                    if row2[i] is not None:
                        header[i] = row2[i]
                data_start = 2
            header = [str(h).strip() if h is not None else f"col{i}"
                      for i, h in enumerate(header)]
            sheets[ws.title] = [dict(zip(header, r)) for r in rows[data_start:]]
        main_rows = []
        for prefer in ["成绩总表", "成绩", "总评"]:
            if prefer in sheets and sheets[prefer]:
                main_rows = sheets[prefer]
                break
        if not main_rows:
            main_rows = next(iter(sheets.values()), [])
        return {"main_rows": main_rows, "sheets": sheets}
    with open(p, "r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    return {"main_rows": rows, "sheets": {"CSV": rows}}


def filter_excluded_docs(docs: list, course_type: str) -> tuple:
    """按业务边界过滤不处理材料（见 business-rules.md 第二节）。
    返回 (保留docs, 排除材料名列表)。"""
    excluded = EXCLUDED_MATERIALS_T1 if course_type == "T1" else set()
    kept, excluded_names = [], []
    for doc in docs:
        name = doc.get("name", "")
        if any(ex in name for ex in excluded):
            excluded_names.append(name)
        else:
            kept.append(doc)
    return kept, excluded_names


def generate_makeup_list(data_rows: list, output_path: Path, course: dict) -> dict:
    """生成补考名单（仅 T1，见 business-rules.md 第一节）。
    有资格：参加考试(卷面≠0且未标记缺考)且综合<60；
    无资格：缺考；待确认：违纪/免考/缓考/特殊标记。"""
    import openpyxl
    from openpyxl.styles import Font
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "补考名单"
    ws.append(["学号", "姓名", "综合成绩", "卷面成绩", "补考资格", "备注"])
    stats = {"挂科数": 0, "有资格": 0, "无资格": 0, "待确认": 0, "待确认学生": []}
    for r in data_rows:
        total = r.get("最终成绩") or r.get("总评") or r.get("综合成绩") or ""
        exam = r.get("考试成绩") or r.get("期末") or r.get("卷面") or ""
        sid = r.get("学号", "")
        sname = r.get("姓名", "")
        remark = str(r.get("备注", "")) + str(r.get("状态", ""))
        try:
            total_f = float(total)
        except (ValueError, TypeError):
            continue
        if total_f >= 60:
            continue
        stats["挂科数"] += 1
        try:
            exam_f = float(exam) if exam != "" else 0
        except (ValueError, TypeError):
            exam_f = 0
        if "缺考" in remark:
            status = "否（缺考）"
            stats["无资格"] += 1
        elif any(k in remark for k in ["违纪", "免考", "缓考", "特殊"]):
            status = "待确认"
            stats["待确认"] += 1
            stats["待确认学生"].append((sid, sname, remark))
        else:
            status = "是"
            stats["有资格"] += 1
        ws.append([sid, sname, total, exam, status, remark])
    for cell in ws[1]:
        cell.font = Font(bold=True)
    ws.cell(row=ws.max_row + 2, column=1, value=AI_NOTE)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    return stats


# --------------------------------------------------------------------------- #
# 过程性考核分项：根据大纲分析构成，从教师数据提取对应分项
# 详见 references/business-rules.md 第七节
# --------------------------------------------------------------------------- #
ASSESSMENT_SYNONYMS = {
    "课堂表现": ["课堂表现", "考勤", "课堂"],
    "平时作业": ["平时作业", "作业"],
    "上机操作": ["上机操作", "上机", "实验", "实验成绩"],
    "实验": ["实验", "上机"],
    "考勤": ["考勤", "课堂表现"],
    "课堂讨论": ["课堂讨论", "讨论", "互动"],
    "阶段测试": ["阶段测试", "测试", "测验"],
    "课程报告": ["课程报告", "报告", "论文"],
}


def _match_link_to_column(link: str, headers: list):
    """将大纲环节名匹配到数据列名（含同义词）。"""
    syns = ASSESSMENT_SYNONYMS.get(link, [link])
    for h in headers:
        for kw in syns:
            if kw in str(h):
                return h
    return None


def match_assessment_to_data(composition: dict, data_bundle: dict) -> list:
    """将大纲过程性考核分项匹配到教师数据中的对应列。
    匹配顺序：主表列名 → 其他sheet列名。返回 [{环节,权重,权重值,细则,数据列,数据来源}]。"""
    sheets = data_bundle.get("sheets", {})
    main_rows = data_bundle.get("main_rows", [])
    main_header = list(main_rows[0].keys()) if main_rows else []
    matched = []
    for item in composition.get("process", []):
        link = item["环节"]
        col = _match_link_to_column(link, main_header)
        source = "成绩总表" if col else None
        if not col:
            for sname, srows in sheets.items():
                if sname == "成绩总表" or not srows:
                    continue
                sheader = list(srows[0].keys()) if srows else []
                col = _match_link_to_column(link, sheader)
                if col:
                    source = sname
                    break
        matched.append({**item, "数据列": col, "数据来源": source})
    return matched


def gen_process_score_sheet(dst: Path, data_bundle: dict, matched: list, course: dict) -> dict:
    """生成过程考核成绩登记表：按大纲分项组织，从教师数据提取对应分项。
    详见 references/business-rules.md 第七节。"""
    import openpyxl
    from openpyxl.styles import Font
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "过程考核成绩登记表"
    headers = ["学号", "姓名"] + [
        f"{m['环节']}({m['权重']})" for m in matched] + ["过程成绩合计"]
    ws.append(headers)

    main_rows = data_bundle.get("main_rows", [])
    sheets = data_bundle.get("sheets", {})
    # 每个分项建立 学号->分数 映射
    col_maps = []
    for m in matched:
        col_map = {}
        source = m.get("数据来源")
        col = m.get("数据列")
        if source and col and source in sheets:
            for r in sheets[source]:
                sid = str(r.get("学号", ""))
                try:
                    col_map[sid] = float(r.get(col, 0) or 0)
                except (ValueError, TypeError):
                    col_map[sid] = ""
        col_maps.append(col_map)

    filled = 0
    for r in main_rows:
        sid = str(r.get("学号", ""))
        sname = r.get("姓名", "")
        row_data = [sid, sname]
        total = 0.0
        total_w = 0.0
        for i, m in enumerate(matched):
            val = col_maps[i].get(sid, "")
            row_data.append(val)
            try:
                w = float(m.get("权重值", 0))
                total += float(val) * w / 100
                total_w += w
            except (ValueError, TypeError):
                pass
        row_data.append(round(total, 1) if total_w > 0 else "")
        ws.append(row_data)
        filled += 1

    for cell in ws[1]:
        cell.font = Font(bold=True)
    ws.cell(row=ws.max_row + 2, column=1, value=AI_NOTE)
    dst.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(dst))
    return {"学生数": filled, "分项数": len(matched),
            "未匹配分项": [m["环节"] for m in matched if not m.get("数据列")]}


# --------------------------------------------------------------------------- #
# 分析报告条件生成 + 归档目录填写（详见 business-rules.md 第八、九节）
# --------------------------------------------------------------------------- #
def detect_analysis_materials(materials_dir) -> dict:
    """检测用户材料中的试卷/参考答案/考核方案（见 business-rules.md 第八节）。
    返回 {'试卷':[], '参考答案':[], '考核方案':[], 'has_any': bool, 'has_all': bool}。"""
    result = {"试卷": [], "参考答案": [], "考核方案": [], "has_any": False, "has_all": False}
    if not materials_dir:
        return result
    d = Path(materials_dir)
    if not d.exists():
        return result
    files = [f for f in d.rglob("*") if f.is_file() and not is_image_or_temp(f)]
    for f in files:
        name = f.name
        if any(k in name for k in ["试卷", "样卷"]) and "答案" not in name and "分析" not in name:
            result["试卷"].append(str(f))
        if any(k in name for k in ["参考答案", "答案", "评分细则", "评分标准"]) and "分析" not in name:
            result["参考答案"].append(str(f))
        if "考核方案" in name or "课程考核方案" in name:
            result["考核方案"].append(str(f))
    result["has_any"] = bool(result["试卷"] or result["参考答案"] or result["考核方案"])
    result["has_all"] = bool(result["试卷"] and result["参考答案"] and result["考核方案"])
    return result


def generate_analysis_report(template_path, dst: Path, values: dict,
                             analysis_materials: dict, course: dict, data_rows: list):
    """生成分析报告（见 business-rules.md 第八节）。
    以高校教学专家角度，根据用户数据（试卷/参考答案/考核方案+成绩）生成。
    返回 (status, note)。"""
    # 1. 无素材 → 提示无法编写
    if not analysis_materials.get("has_any"):
        note_path = dst.parent / ("无法编写_" + dst.stem + ".txt")
        dst.parent.mkdir(parents=True, exist_ok=True)
        note_path.write_text(
            f"【无法编写】分析报告\n\n用户数据中未提供试卷/参考答案/考核方案，"
            f"无法以高校教学专家角度编写分析报告。\n请提供上述材料后重新生成。\n\n{AI_NOTE}\n",
            encoding="utf-8")
        return "无法编写", "用户数据缺少试卷/参考答案/考核方案（见 business-rules.md 第八节）"
    # 2. 无模板 → 提示模板缺失
    if not template_path or not template_path.exists():
        note_path = dst.parent / ("模板缺失_" + dst.stem + ".txt")
        dst.parent.mkdir(parents=True, exist_ok=True)
        note_path.write_text(
            f"【模板缺失】分析报告\n\n模板目录中未找到分析报告模板，"
            f"已检测到用户素材但无法生成。\n请在模板目录中补充分析报告模板（附件名含'分析'）。\n\n{AI_NOTE}\n",
            encoding="utf-8")
        return "模板缺失", "模板目录中未找到分析报告模板"
    # 3. .doc 老格式 → 提示需转换
    ext = template_path.suffix.lower()
    if ext == ".doc":
        note_path = dst.parent / ("待转换_" + dst.stem + ".txt")
        dst.parent.mkdir(parents=True, exist_ok=True)
        note_path.write_text(
            f"【模板需转换】分析报告\n\n分析报告模板 {template_path.name} 为 .doc 老格式，"
            f"python-docx 无法读取。\n请转换为 .docx 后重新生成，或手动填写。\n\n"
            f"已检测到用户素材：试卷 {len(analysis_materials['试卷'])} 份、"
            f"参考答案 {len(analysis_materials['参考答案'])} 份、"
            f"考核方案 {len(analysis_materials['考核方案'])} 份。\n\n{AI_NOTE}\n",
            encoding="utf-8")
        return "模板需转换", f"模板为.doc老格式: {template_path.name}"
    # 4. .docx → 填模板 + 附加专家分析
    if ext == ".docx":
        try:
            missed = fill_docx(template_path, dst, values)
            # 附加基于成绩与素材的专家分析
            from docx import Document
            doc = Document(str(dst))
            doc.add_paragraph("")
            doc.add_heading("AI 高校教学专家分析（基于用户数据）", level=2)
            if data_rows:
                scores = []
                for r in data_rows:
                    v = r.get("最终成绩") or r.get("总评") or r.get("综合成绩")
                    try:
                        scores.append(float(v))
                    except (ValueError, TypeError):
                        pass
                if scores:
                    n = len(scores)
                    mean = sum(scores) / n
                    passed = sum(1 for s in scores if s >= 60)
                    excellent = sum(1 for s in scores if s >= 85)
                    doc.add_paragraph(
                        f"参评人数 {n}；平均分 {mean:.1f}；及格率 {passed/n*100:.1f}%；"
                        f"优秀率 {excellent/n*100:.1f}%。")
            doc.add_paragraph(
                f"分析依据素材：试卷 {len(analysis_materials['试卷'])} 份、"
                f"参考答案 {len(analysis_materials['参考答案'])} 份、"
                f"考核方案 {len(analysis_materials['考核方案'])} 份。")
            doc.add_paragraph("（详细题目设计/覆盖度/难度分析请结合试卷内容完善。）")
            doc.save(str(dst))
            note = "未命中占位符: " + ", ".join(sorted(set(missed))) if missed else ""
            return "已生成(分析报告)", note
        except Exception as e:
            return "失败", str(e)
    return "跳过", f"分析报告模板格式不支持: {ext}"


def fill_archive_directory(template_path, dst: Path, manifest: list,
                            course: dict, values: dict):
    """填写归档目录：填课程信息占位符 + 根据材料清单填写表格（见 business-rules.md 第九节）。
    返回 (status, note)。"""
    if not template_path or not template_path.exists():
        _placeholder(dst.parent, dst.name, "归档目录")
        return "待补充", "未找到归档目录模板"
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(template_path))
        ws = wb.active
        from openpyxl.cell.cell import MergedCell
        # 1. 填充占位符（课程信息），跳过合并单元格的只读单元格
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell, MergedCell):
                    continue
                if isinstance(cell.value, str):
                    cell.value, _ = replace_text(cell.value, values)
        # 2. 找材料清单表头（含"序号"+"材料名称"或"归档材料"）
        rows = list(ws.iter_rows(values_only=True))
        header_idx, cols = -1, {}
        for i, row in enumerate(rows):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if "序号" in cells and ("材料名称" in cells or "归档材料" in cells):
                header_idx = i
                cols["序号"] = cells.index("序号")
                cols["材料名称"] = cells.index("材料名称") if "材料名称" in cells else cells.index("归档材料")
                cols["份数"] = cells.index("份数") if "份数" in cells else -1
                cols["备注"] = cells.index("备注") if "备注" in cells else -1
                break
        # 3. 在表头行后填写材料清单
        if header_idx >= 0:
            seq = 1
            for folder, name, fname, status, note in manifest:
                if status == "不处理":
                    continue
                r = header_idx + 1 + seq
                for col_key, val in [("序号", seq), ("材料名称", name)]:
                    if col_key in cols:
                        try:
                            ws.cell(row=r, column=cols[col_key] + 1, value=val)
                        except AttributeError:
                            pass  # 跳过合并单元格
                if cols.get("份数", -1) >= 0:
                    try:
                        ws.cell(row=r, column=cols["份数"] + 1, value=1)
                    except AttributeError:
                        pass
                if cols.get("备注", -1) >= 0:
                    try:
                        ws.cell(row=r, column=cols["备注"] + 1, value=status)
                    except AttributeError:
                        pass
                seq += 1
        ws.cell(row=ws.max_row + 2, column=1, value=AI_NOTE)
        dst.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(dst))
        filled = sum(1 for m in manifest if m[3] != "不处理")
        return f"已生成(归档目录{filled}项)", ""
    except Exception as e:
        return "失败", str(e)


# --------------------------------------------------------------------------- #
# 占位符替换
# --------------------------------------------------------------------------- #
def build_values(course: dict, type_info: dict) -> dict:
    """构造占位符替换表，补齐派生字段。"""
    vals = dict(course)
    vals.setdefault("日期", _dt.date.today().isoformat())
    sem = vals.get("学期", "")
    m = re.match(r"(\d{4}-\d{4})", sem)
    if m and "学年" not in vals:
        vals["学年"] = m.group(1)
    vals["课程类型"] = type_info.get("course_type_name", "")
    vals["课程类型代码"] = type_info.get("course_type_code", "")
    vals["是否含课内实验"] = "是" if type_info.get("need_lab_archive") else "否"
    return vals


def replace_text(text: str, values: dict) -> tuple:
    if not isinstance(text, str):
        return text, []
    missed = []

    def _sub(m):
        key = m.group(1).strip()
        if key in values and values[key] not in (None, ""):
            return str(values[key])
        missed.append(key)
        return m.group(0)

    return PLACEHOLDER_RE.sub(_sub, text), missed


# --------------------------------------------------------------------------- #
# 模板匹配
# --------------------------------------------------------------------------- #
def match_template_dir(templates_root: Path, type_info: dict) -> Path:
    """在模板根目录下匹配课程类型对应的子文件夹。
    返回匹配的子文件夹路径；若无子文件夹则返回根目录本身；若多候选返回 None（需用户确认）。"""
    if not templates_root.exists():
        return templates_root
    subdirs = [d for d in templates_root.iterdir() if d.is_dir()]
    if not subdirs:
        return templates_root  # 无子文件夹，用根目录

    folder_kws = type_info.get("folder_keywords", [])
    type_code = type_info.get("course_type_code", "")
    type_name = type_info.get("course_type_name", "")

    candidates = []
    for d in subdirs:
        name = d.name
        # 精确匹配：子文件夹名同时含"理论课"+"考试/考查"（针对 T1/T2 区分）
        if type_code == "theory_exam" and "考试" in name and "理论课" in name:
            candidates.append(d)
        elif type_code == "theory_assess" and "考查" in name and "理论课" in name:
            candidates.append(d)
        elif type_code in ("course_design", "training", "internship"):
            # T3/T4/T5：子文件夹名含任一关键词
            if any(kw in name for kw in folder_kws):
                candidates.append(d)
        else:
            # 兜底：含任一关键词
            if any(kw in name for kw in folder_kws):
                candidates.append(d)

    if len(candidates) == 1:
        return candidates[0]
    if len(candidates) > 1:
        # 多候选，返回第一个但记录警告（build 流程会提示）
        return candidates[0]
    # 未匹配，返回根目录（让流程继续，但模板可能不全）
    return templates_root


def find_template(template_dir: Path, template_name: str) -> Path:
    """在模板目录下查找指定模板文件（含子目录）。"""
    # 先精确匹配
    p = template_dir / template_name
    if p.exists():
        return p
    # 模糊匹配（文件名包含 template_name 的 stem）
    stem = Path(template_name).stem
    for fp in template_dir.rglob("*"):
        if fp.is_file() and stem in fp.stem:
            return fp
    return p  # 返回不存在的路径，由调用方处理


# --------------------------------------------------------------------------- #
# 文档生成
# --------------------------------------------------------------------------- #
def fill_docx(src: Path, dst: Path, values: dict) -> list:
    from docx import Document
    doc = Document(str(src))
    missed = []

    def _runs(paragraph):
        for run in paragraph.runs:
            run.text, m = replace_text(run.text, values)
            missed.extend(m)

    for p in doc.paragraphs:
        _runs(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _runs(p)
    # 追加 AI 附注
    doc.add_paragraph("")
    doc.add_paragraph(AI_NOTE)
    doc.save(str(dst))
    return missed


def fill_xlsx(src: Path, dst: Path, values: dict) -> list:
    import openpyxl
    wb = openpyxl.load_workbook(str(src))
    missed = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value, m = replace_text(cell.value, values)
                    missed.extend(m)
        # 在最后一个 sheet 末尾追加附注
    # 追加附注到 active sheet
    ws = wb.active
    max_row = ws.max_row or 1
    ws.cell(row=max_row + 2, column=1, value=AI_NOTE)
    wb.save(str(dst))
    return missed


def copy_file(src: Path, dst: Path) -> None:
    shutil.copy2(str(src), str(dst))


def gen_data_xlsx(dst: Path, rows: list, weights: dict, title: str = "成绩") -> None:
    import openpyxl
    from openpyxl.styles import Font
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title
    if not rows:
        ws["A1"] = "（无数据）"
        ws.cell(row=3, column=1, value=AI_NOTE)
        wb.save(str(dst))
        return
    headers = list(rows[0].keys())
    if "总评" not in headers and "平时" in headers and "期末" in headers:
        w_p = float(weights.get("平时", 0)) / 100
        w_m = float(weights.get("期中", 0)) / 100
        w_f = float(weights.get("期末", 0)) / 100
        total_w = w_p + w_m + w_f or 1
        headers = headers + ["总评"]
        for r in rows:
            try:
                p = float(r.get("平时", 0) or 0)
                m = float(r.get("期中", 0) or 0)
                f = float(r.get("期末", 0) or 0)
                r["总评"] = round((p * w_p + m * w_m + f * w_f) / total_w, 1)
            except (ValueError, TypeError):
                r["总评"] = ""
    ws.append(headers)
    for r in rows:
        ws.append([r.get(h, "") for h in headers])
    for cell in ws[1]:
        cell.font = Font(bold=True)
    # 追加附注
    ws.cell(row=ws.max_row + 2, column=1, value=AI_NOTE)
    wb.save(str(dst))


def _placeholder(dest_dir: Path, fname: str, name: str):
    note = dest_dir / ("待补充_" + Path(fname).stem + ".txt")
    note.write_text(
        f"【待补充】{name}\n\n本材料尚未提供模板或数据，请事后手动补充后放入此文件夹。\n\n{AI_NOTE}\n",
        encoding="utf-8",
    )


# --------------------------------------------------------------------------- #
# 数据读取
# --------------------------------------------------------------------------- #
def read_data(path: str) -> list:
    if not path:
        return []
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"平时数据文件不存在: {path}")
    if p.suffix.lower() in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return []
        header = [str(h).strip() if h is not None else f"col{i}" for i, h in enumerate(rows[0])]
        return [dict(zip(header, r)) for r in rows[1:]]
    with open(p, "r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


# --------------------------------------------------------------------------- #
# 主流程
# --------------------------------------------------------------------------- #
def resolve_filename(pattern: str, values: dict) -> str:
    name, _ = replace_text(pattern, values)
    return name


def build(cfg: dict):
    course = cfg.get("course", {})
    syllabus_file = cfg.get("syllabus_file")
    data_file = cfg.get("data_file")
    templates_dir = Path(cfg.get("templates_dir") or DEFAULT_TEMPLATES)
    output_dir = Path(cfg.get("output_dir") or ".")
    weights = cfg.get("weights", {"平时": 40, "期中": 0, "期末": 60})
    materials_dir = cfg.get("materials_dir") or (str(Path(data_file).parent) if data_file else "")

    # 1. 识别课程类型
    type_info = None
    if syllabus_file:
        print(f"📖 识别课程类型: {syllabus_file}")
        type_info = ict.identify(syllabus_file)
        print(f"   → 课程类型: {type_info['course_type_name']}（{type_info['course_type']}）")
        print(f"   → 是否含课内实验: {'是' if type_info['need_lab_archive'] else '否'}")
    else:
        # 无大纲，默认 T1，提示用户提供
        print("⚠️ 未提供大纲文件，默认按 T1 理论课（考试）处理，建议提供大纲以精确识别。")
        type_info = {
            "course_type": "T1", "course_type_code": "theory_exam",
            "course_type_name": "理论课（考试）", "folder_keywords": ["理论课", "考试"],
            "need_lab_archive": False, "evidence": {"note": "未提供大纲，默认 T1"},
        }

    course_type = type_info["course_type"]
    need_lab = type_info["need_lab_archive"]

    # 2. 动态发现模板（不写死，见 template-discovery.md）
    tpl = discover_templates(templates_dir, type_info)
    print(f"📂 主模板目录: {tpl['matched_dir'] or templates_dir}")
    if tpl['checklist_path']:
        print(f"📋 归档确认单: {tpl['checklist_path'].name}")
    if tpl['requirement_path']:
        print(f"📄 要求文件: {tpl['requirement_path'].name}")
    print(f"📎 附件模板: {len(tpl['attachments'])} 个（图片示例 {len(tpl['images'])} 个已跳过）")
    template_dir = tpl['matched_dir'] or templates_dir

    # 3. 加载归档规范（DEFAULT_CONVENTION 作材料清单框架，实际模板动态发现）
    convention = DEFAULT_CONVENTION.get(course_type, DEFAULT_CONVENTION["T1"])
    values = build_values(course, type_info)

    # 4. 读取数据（多 Sheet，见 business-rules.md 第五节）
    data_bundle = read_data_multi_sheet(data_file)
    data_rows = data_bundle["main_rows"]
    if data_file and not data_rows:
        print(f"⚠️ 数据文件读取为空: {data_file}")
    elif data_rows:
        print(f"📊 成绩数据: {len(data_rows)} 行，sheets={list(data_bundle['sheets'].keys())}")

    # 4.5 解析大纲考核构成（过程性分项，供过程考核成绩登记表生成，见 business-rules.md 第七节）
    matched_assessment = []
    if syllabus_file:
        composition = ict.parse_assessment_composition(syllabus_file)
        matched_assessment = match_assessment_to_data(composition, data_bundle)
        if matched_assessment:
            parts = [f"{m['环节']}({m['权重']}→{m.get('数据来源') or '未匹配'}/{m.get('数据列') or '无'})"
                     for m in matched_assessment]
            print(f"📐 过程性考核分项({len(matched_assessment)}项): " + " | ".join(parts))
            unmatched = [m['环节'] for m in matched_assessment if not m.get('数据列')]
            if unmatched:
                print(f"   ⚠️ 未匹配到数据的分项: {', '.join(unmatched)}（需向用户确认数据列）")

    # 4.6 检测分析报告素材（试卷/参考答案/考核方案，见 business-rules.md 第八节）
    analysis_materials = detect_analysis_materials(materials_dir)
    if analysis_materials["has_any"]:
        print(f"📐 分析报告素材: 试卷{len(analysis_materials['试卷'])}份、"
              f"参考答案{len(analysis_materials['参考答案'])}份、"
              f"考核方案{len(analysis_materials['考核方案'])}份")
    else:
        print("📐 分析报告素材: 未检测到试卷/参考答案/考核方案（分析报告将提示无法编写）")

    # 5. 生成归档目录
    root_name = resolve_filename(
        convention.get("root_folder", "{{课程名}}_{{学期}}_课程归档"), values)
    # 规范中没有 root_folder，用默认
    if "{{" in root_name:
        root_name = f"{course.get('课程名', '未命名')}_{course.get('学期', '')}_课程归档"
    root_path = output_dir / root_name
    root_path.mkdir(parents=True, exist_ok=True)

    manifest = []  # (folder, name, fname, status, note)
    all_excluded = []  # 不处理材料（由教师/教务秘书处理，见 business-rules.md 第二节）

    # 6. 生成主归档材料（按业务边界过滤不处理材料）
    for cat in convention.get("categories", []):
        folder = cat["folder"]
        dest_dir = root_path / folder
        dest_dir.mkdir(parents=True, exist_ok=True)
        docs, excl = filter_excluded_docs(cat.get("docs", []), course_type)
        for ex in excl:
            all_excluded.append((folder, ex))
        for doc in docs:
            name = doc["name"]
            fname = resolve_filename(doc["file"], values)
            dst = dest_dir / fname
            kind = doc.get("kind", "template")
            status, note = "已生成", ""

            try:
                if kind == "copy":
                    tmpl = find_template(template_dir, doc.get("template", ""))
                    if tmpl.exists():
                        copy_file(tmpl, dst)
                        status = "已复制"
                    else:
                        _placeholder(dest_dir, fname, name)
                        status, note = "待补充", f"缺少模板: {doc.get('template')}"
                elif kind == "template":
                    tmpl = find_template(template_dir, doc.get("template", ""))
                    if tmpl.exists():
                        tmpl_name = tmpl.name
                        if "分析" in tmpl_name:
                            # 分析报告：条件生成（见 business-rules.md 第八节）
                            status, note = generate_analysis_report(
                                tmpl, dst, values, analysis_materials, course, data_rows)
                        elif "归档目录" in tmpl_name:
                            # 归档目录：根据材料清单填写（见 business-rules.md 第九节）
                            status, note = fill_archive_directory(
                                tmpl, dst, manifest, course, values)
                        else:
                            ext = tmpl.suffix.lower()
                            if ext == ".docx":
                                missed = fill_docx(tmpl, dst, values)
                            elif ext in (".xlsx", ".xlsm"):
                                missed = fill_xlsx(tmpl, dst, values)
                            else:
                                txt = tmpl.read_text(encoding="utf-8", errors="ignore")
                                txt, missed = replace_text(txt, values)
                                txt += f"\n\n{AI_NOTE}\n"
                                dst.write_text(txt, encoding="utf-8")
                            status = "已生成(模板填充)"
                            if missed:
                                note = "未命中占位符: " + ", ".join(sorted(set(missed)))
                    else:
                        _placeholder(dest_dir, fname, name)
                        status, note = "待补充", f"缺少模板: {doc.get('template')}"
                elif kind == "data":
                    if not data_rows:
                        _placeholder(dest_dir, fname, name)
                        status, note = "待补充", "未提供平时数据"
                    elif ("过程考核" in name or "平时成绩" in name or "过程性" in name) and matched_assessment:
                        # 过程考核成绩登记表：按大纲分项组织（见 business-rules.md 第七节）
                        info = gen_process_score_sheet(dst, data_bundle, matched_assessment, course)
                        status = f"已生成(过程性分项{info['分项数']}项,学生{info['学生数']}人)"
                        if info["未匹配分项"]:
                            note = "未匹配分项: " + ", ".join(info["未匹配分项"])
                    else:
                        gen_data_xlsx(dst, data_rows, weights, name)
                        status = "已生成(数据表)"
            except Exception as e:
                status, note = "失败", str(e)

            manifest.append((folder, name, fname, status, note))

    # 7. 课内实验附加材料
    if need_lab:
        lab_dir = root_path / "05_课内实验"
        lab_dir.mkdir(parents=True, exist_ok=True)
        lab_template_dir = templates_dir / "课内实验"
        if not lab_template_dir.exists():
            # 模糊匹配
            for d in templates_dir.iterdir() if templates_dir.exists() else []:
                if d.is_dir() and any(kw in d.name for kw in ["课内实验", "实验", "上机", "实践"]):
                    lab_template_dir = d
                    break
        for doc in LAB_CONVENTION["categories"][0]["docs"]:
            name = doc["name"]
            fname = resolve_filename(doc["file"], values)
            dst = lab_dir / fname
            try:
                tmpl = find_template(lab_template_dir, doc.get("template", ""))
                if tmpl.exists():
                    ext = tmpl.suffix.lower()
                    if ext == ".docx":
                        missed = fill_docx(tmpl, dst, values)
                    elif ext in (".xlsx", ".xlsm"):
                        missed = fill_xlsx(tmpl, dst, values)
                    else:
                        txt = tmpl.read_text(encoding="utf-8", errors="ignore")
                        txt, missed = replace_text(txt, values)
                        txt += f"\n\n{AI_NOTE}\n"
                        dst.write_text(txt, encoding="utf-8")
                    status = "已生成(课内实验)"
                    note = "未命中占位符: " + ", ".join(sorted(set(missed))) if missed else ""
                else:
                    _placeholder(lab_dir, fname, name)
                    status, note = "待补充", f"缺少模板: {doc.get('template')}"
            except Exception as e:
                status, note = "失败", str(e)
            manifest.append(("05_课内实验", name, fname, status, note))

    # 7.4 归档目录动态发现与填写（见 business-rules.md 第九节）
    # 若循环中未生成归档目录（DEFAULT_CONVENTION 未列），从模板附件动态发现并填写
    has_archive_dir = any("归档目录" in m[1] for m in manifest)
    if not has_archive_dir and tpl.get('attachments'):
        for att in tpl['attachments']:
            if "归档目录" in att.name and att.suffix.lower() == ".xlsx":
                ad_name = att.stem.split("：")[-1] if "：" in att.stem else att.stem
                ad_fname = f"{ad_name}_{course.get('课程名','')}_{course.get('学期','')}.xlsx"
                ad_dst = root_path / ad_fname
                ad_status, ad_note = fill_archive_directory(att, ad_dst, manifest, course, values)
                manifest.append(("归档目录", ad_name, ad_fname, ad_status, ad_note))
                print(f"📋 归档目录: {ad_fname} → {ad_status}")
                break

    # 7.5 补考名单（仅 T1，见 business-rules.md 第一节）
    if course_type == "T1" and data_rows:
        makeup_name = f"补考名单_{course.get('课程名','')}_{course.get('学期','')}.xlsx"
        makeup_path = root_path / makeup_name
        print("📝 生成补考名单（仅 T1）...")
        makeup_stats = generate_makeup_list(data_rows, makeup_path, course)
        manifest.append(("补考", "补考名单", makeup_name,
                         f"已生成(挂科{makeup_stats['挂科数']}·有资格{makeup_stats['有资格']}·待确认{makeup_stats['待确认']})", ""))
        if makeup_stats["待确认学生"]:
            print(f"   ⚠️ {len(makeup_stats['待确认学生'])} 名学生补考资格待确认（见 business-rules.md 1.3）：")
            for sid, sname, remark in makeup_stats["待确认学生"][:5]:
                print(f"      {sid} {sname}（{remark}）")
            if len(makeup_stats["待确认学生"]) > 5:
                print(f"      ...(共{len(makeup_stats['待确认学生'])}名)")

    # 7.6 不处理材料提示（见 business-rules.md 第二节）
    if all_excluded:
        print(f"\n⚠️ 以下 {len(all_excluded)} 项材料不属于技能处理范围（由教师/教务秘书处理）：")
        for folder, name in all_excluded:
            print(f"   - [{folder}] {name}")
            manifest.append((folder, name, "（不处理）", "不处理", "由教师/教务秘书处理，见 business-rules.md"))

    # 8. 写归档清单
    _write_manifest(root_path, root_name, course, type_info, manifest)

    # 9. 规范性校验（传入归档确认单作为校验依据，见 validation-rules.md 3.5）
    print("🔍 执行规范性校验...")
    checklist_items = parse_archive_checklist(tpl['checklist_path']) if tpl['checklist_path'] else []
    report_text, _ = va.validate(root_path, course, data_rows, course_type, need_lab, weights,
                                  checklist_items=checklist_items, data_bundle=data_bundle)
    (root_path / "规范性校验报告.md").write_text(report_text, encoding="utf-8")

    # 10. 教学专家分析
    print("🎓 生成教学专家分析报告...")
    expert_text, _ = er.generate(course, data_rows, type_info["course_type_name"])
    (root_path / "教学专家分析报告.md").write_text(expert_text, encoding="utf-8")

    # 11. 打包 ZIP
    zip_path = output_dir / (root_name + ".zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in sorted(root_path.rglob("*")):
            if fp.is_file():
                zf.write(fp, fp.relative_to(output_dir))

    print(f"\n✅ 归档已生成: {root_path}")
    print(f"✅ 压缩包: {zip_path}")
    done = sum(1 for m in manifest if m[3] not in ("待补充", "失败"))
    print(f"📋 共 {len(manifest)} 项材料，成功 {done} 项，待补充/失败 {len(manifest) - done} 项。")
    print(f"📄 规范性校验报告: {root_path / '规范性校验报告.md'}")
    print(f"📄 教学专家分析报告: {root_path / '教学专家分析报告.md'}")


def _write_manifest(root_path: Path, root_name: str, course: dict, type_info: dict, manifest: list):
    md = root_path / "归档清单.md"
    csvp = root_path / "归档清单.csv"
    lines = [f"# 归档清单 — {root_name}", "", "## 课程信息"]
    for k, v in course.items():
        lines.append(f"- {k}: {v}")
    lines.append(f"- 课程类型: {type_info.get('course_type_name', '')}（{type_info.get('course_type', '')}）")
    lines.append(f"- 是否含课内实验: {'是' if type_info.get('need_lab_archive') else '否'}")
    lines += ["", "## 材料明细", "",
              "| 类别 | 材料 | 文件名 | 状态 | 说明 |",
              "|------|------|--------|------|------|"]
    for folder, name, fname, status, note in manifest:
        lines.append(f"| {folder} | {name} | {fname} | {status} | {note} |")
    lines += ["", "## 附加报告",
              "- 规范性校验报告.md",
              "- 教学专家分析报告.md",
              "",
              "---",
              "",
              AI_NOTE]
    md.write_text("\n".join(lines), encoding="utf-8")

    with open(csvp, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["类别", "材料", "文件名", "状态", "说明"])
        for row in manifest:
            w.writerow(row)


# --------------------------------------------------------------------------- #
# 命令行入口
# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser(description="lzhit-teaching-archive 高校期末教学档案归档生成")
    ap.add_argument("--config", help="JSON 配置文件路径（推荐）")
    ap.add_argument("--course-name", help="课程名")
    ap.add_argument("--course-code", help="课程代码")
    ap.add_argument("--semester", help="学期，如 2025-2026-2")
    ap.add_argument("--teacher", help="任课教师")
    ap.add_argument("--class", dest="klass", help="班级")
    ap.add_argument("--syllabus", help="课程教学大纲文件路径（docx/txt/md/pdf）")
    ap.add_argument("--data-file", help="平时数据 CSV/xlsx 路径")
    ap.add_argument("--templates-dir", help="模板根目录（默认技能内 assets/templates）")
    ap.add_argument("--output-dir", help="输出根目录")
    args = ap.parse_args()

    if args.config:
        with open(args.config, "r", encoding="utf-8") as f:
            cfg = json.load(f)
    else:
        cfg = {
            "course": {
                "课程名": args.course_name or "未命名课程",
                "课程代码": args.course_code or "",
                "学期": args.semester or "",
                "教师": args.teacher or "",
                "班级": args.klass or "",
            },
            "syllabus_file": args.syllabus,
            "data_file": args.data_file,
            "templates_dir": args.templates_dir,
            "output_dir": args.output_dir,
        }

    build(cfg)


if __name__ == "__main__":
    main()
