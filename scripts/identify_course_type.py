#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
课程类型识别脚本（lzhit-teaching-archive）

读取课程教学大纲，识别课程类型（T1-T5）及是否需要课内实验归档（L1）。
识别规则见 references/course-type-rules.md。

用法:
    python identify_course_type.py --syllabus path/to/syllabus.docx
    python identify_course_type.py --syllabus syllabus.docx --json

依赖: python-docx（读取 .docx 大纲；.txt/.md 无需依赖）
"""

import argparse
import json
import sys
from pathlib import Path

# --------------------------------------------------------------------------- #
# 关键词定义（按优先级 T5 > T4 > T3 > T1 > T2）
# --------------------------------------------------------------------------- #
# T2 去掉了过宽泛的"报告""论文"，避免在 T1 大纲中误命中
KEYWORDS = {
    "T5": ["实习", "实习报告", "实习鉴定", "毕业实习", "专业实习", "顶岗实习"],
    "T4": ["实训", "实训报告", "操作考核", "技能考核", "集中实训"],
    "T3": ["课程设计", "设计任务", "设计报告", "工程设计", "综合设计"],
    "T1": ["考试", "闭卷考试", "开卷考试", "机试", "笔试", "期末考试", "统一考试"],
    "T2": ["大作业", "课程报告", "课程论文", "答辩", "作品", "综合报告", "调研报告"],
}

LAB_KEYWORDS = ["课内实验", "上机", "实践", "实验学时", "上机学时", "实践学时", "实验环节"]

TYPE_INFO = {
    "T1": {"code": "theory_exam", "name": "理论课（考试）", "folder_keywords": ["理论课", "考试"]},
    "T2": {"code": "theory_assess", "name": "理论课（考查）", "folder_keywords": ["理论课", "考查"]},
    "T3": {"code": "course_design", "name": "课程设计", "folder_keywords": ["课程设计", "设计"]},
    "T4": {"code": "training", "name": "实训", "folder_keywords": ["实训"]},
    "T5": {"code": "internship", "name": "实习", "folder_keywords": ["实习"]},
}

PRIORITY = ["T5", "T4", "T3", "T1", "T2"]


# --------------------------------------------------------------------------- #
# 大纲读取
# --------------------------------------------------------------------------- #
def read_syllabus(path: str) -> str:
    """读取大纲文件，返回全文文本。支持 docx/txt/md/pdf。"""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"大纲文件不存在: {path}")
    ext = p.suffix.lower()
    if ext == ".docx":
        return _read_docx(p)
    if ext == ".pdf":
        return _read_pdf(p)
    return p.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("读取 .docx 大纲需要 python-docx，请安装: pip install python-docx")
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        import PyPDF2  # type: ignore
    except ImportError:
        return "[PDF 文本层提取失败：未安装 PyPDF2。请提供 docx/txt/md 版大纲或安装 PyPDF2]"
    parts = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# 字段提取
# --------------------------------------------------------------------------- #
_NEXT_FIELDS = [
    "考核方式", "考核内容", "成绩构成", "成绩评定", "平时成绩",
    "期末成绩", "教学方式", "课程目标", "教材", "参考书目",
    "课程性质", "课程简介", "教学要求", "教学内容", "学时分配",
]


def extract_field(text: str, field_name: str, max_len: int = 300) -> str:
    """从全文中提取某字段（如"终结性考核"）的内容。
    策略：找字段名位置，取其后到下一个常见字段标题为止。"""
    idx = text.find(field_name)
    if idx < 0:
        return ""
    start = idx + len(field_name)
    end = len(text)
    for nf in _NEXT_FIELDS:
        ni = text.find(nf, start)
        if 0 < ni < end:
            end = ni
    snippet = text[start:end]
    snippet = snippet.lstrip("：: \n\r\t ")
    return snippet[:max_len]


# --------------------------------------------------------------------------- #
# 关键词匹配
# --------------------------------------------------------------------------- #
def find_keywords(text: str, keywords: list) -> list:
    return [kw for kw in keywords if kw in text]


def make_evidence(text: str, keywords: list, field: str) -> dict:
    snippets = []
    for kw in keywords:
        idx = text.find(kw)
        if idx >= 0:
            start = max(0, idx - 20)
            end = min(len(text), idx + len(kw) + 20)
            snippets.append({"keyword": kw, "snippet": text[start:end].replace("\n", " ")})
    return {"matched_keywords": keywords, "source_snippets": snippets, "field": field}


# --------------------------------------------------------------------------- #
# 主识别逻辑
# --------------------------------------------------------------------------- #
def identify(syllabus_path: str) -> dict:
    """识别课程类型。返回结构化 dict。"""
    text = read_syllabus(syllabus_path)

    # 提取"终结性考核"栏；若无则扫全文
    exam_field = extract_field(text, "终结性考核")
    if exam_field:
        search_text = exam_field
        search_field = "终结性考核"
    else:
        # 尝试其他可能字段名
        for alt in ("考核方式", "期末考核", "考核与评价"):
            exam_field = extract_field(text, alt)
            if exam_field:
                search_text = exam_field
                search_field = alt
                break
        else:
            search_text = text
            search_field = "全文（未找到终结性考核栏）"

    course_type = None
    evidence = None
    for type_id in PRIORITY:
        kws = KEYWORDS[type_id]
        hit = find_keywords(search_text, kws)
        if hit:
            course_type = type_id
            evidence = make_evidence(search_text, hit, search_field)
            break

    if course_type is None:
        # 默认 T1，提示用户确认
        course_type = "T1"
        evidence = {
            "matched_keywords": [],
            "source_snippets": [],
            "field": search_field,
            "note": "未匹配任何关键词，默认归为理论课（考试），请用户确认",
        }

    # 课内实验检测（仅 T1/T2）
    need_lab_archive = False
    lab_evidence = None
    if course_type in ("T1", "T2"):
        lab_hit = find_keywords(text, LAB_KEYWORDS)
        if lab_hit:
            need_lab_archive = True
            lab_evidence = make_evidence(text, lab_hit, "全文")

    return {
        "course_type": course_type,
        "course_type_code": TYPE_INFO[course_type]["code"],
        "course_type_name": TYPE_INFO[course_type]["name"],
        "folder_keywords": TYPE_INFO[course_type]["folder_keywords"],
        "need_lab_archive": need_lab_archive,
        "evidence": evidence,
        "lab_evidence": lab_evidence,
        "exam_field_text": exam_field[:200] if exam_field else "",
    }


# --------------------------------------------------------------------------- #
# 考核构成解析（过程性考核分项，供过程考核成绩登记表生成使用）
# 详见 references/business-rules.md 第七节
# --------------------------------------------------------------------------- #
def parse_assessment_composition(syllabus_path: str) -> dict:
    """解析大纲"课程考核与评价"表，返回过程性考核与结果性考核的分项。
    返回 {
        'process': [{'环节', '权重', '细则'}, ...],   # 过程性考核分项
        'result':  [{'环节', '权重', '细则'}, ...],   # 结果性考核分项
        'process_total_weight': int,
        'result_total_weight': int,
    }
    仅支持 .docx（直接读表格）；其他格式返回空。"""
    result = {"process": [], "result": [],
              "process_total_weight": 0, "result_total_weight": 0}
    p = Path(syllabus_path)
    if not p.exists() or p.suffix.lower() != ".docx":
        return result
    try:
        from docx import Document
        doc = Document(str(p))
        for tbl in doc.tables:
            rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
            if not rows:
                continue
            header = rows[0]
            # 识别考核构成表：表头同时含"考核方式"和"考核环节"
            if not (any("考核方式" in h for h in header) and any("考核环节" in h for h in header)):
                continue
            way_idx = next(i for i, h in enumerate(header) if "考核方式" in h)
            link_idx = next(i for i, h in enumerate(header) if "考核环节" in h)
            weight_idx = next((i for i, h in enumerate(header) if "权重" in h), -1)
            detail_idx = next((i for i, h in enumerate(header)
                               if "细则" in h or "评价" in h), -1)
            for row in rows[1:]:
                way = row[way_idx] if way_idx < len(row) else ""
                link = row[link_idx] if link_idx < len(row) else ""
                if not link:
                    continue
                weight = row[weight_idx] if weight_idx >= 0 and weight_idx < len(row) else ""
                detail = row[detail_idx] if detail_idx >= 0 and detail_idx < len(row) else ""
                # 权重提取数字
                import re as _re
                w_nums = _re.findall(r"\d+(?:\.\d+)?", weight)
                w_num = float(w_nums[0]) if w_nums else 0
                item = {"环节": link, "权重": weight, "权重值": w_num, "细则": detail}
                if "过程" in way:
                    result["process"].append(item)
                    result["process_total_weight"] += w_num
                elif "结果" in way or "考试" in way:
                    result["result"].append(item)
                    result["result_total_weight"] += w_num
            break  # 只取第一个匹配的考核构成表
    except Exception:
        pass
    return result


# --------------------------------------------------------------------------- #
# 命令行入口
# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser(description="识别课程类型")
    ap.add_argument("--syllabus", required=True, help="大纲文件路径（docx/txt/md/pdf）")
    ap.add_argument("--json", action="store_true", help="以 JSON 输出")
    args = ap.parse_args()

    try:
        result = identify(args.syllabus)
    except Exception as e:
        print(f"❌ 识别失败: {e}", file=sys.stderr)
        sys.exit(1)

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"课程类型: {result['course_type_name']}（{result['course_type']}）")
        print(f"类型代码: {result['course_type_code']}")
        print(f"是否含课内实验: {'是' if result['need_lab_archive'] else '否'}")
        ev = result.get("evidence", {})
        if ev.get("matched_keywords"):
            print(f"命中关键词: {', '.join(ev['matched_keywords'])}（字段: {ev['field']}）")
        if ev.get("note"):
            print(f"提示: {ev['note']}")
        if result.get("lab_evidence"):
            print(f"课内实验命中: {', '.join(result['lab_evidence']['matched_keywords'])}")


if __name__ == "__main__":
    main()
